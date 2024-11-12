import streamlit as st
import numpy as np
import cv2
from PIL import Image
import pytesseract
import re
from sympy import latex, parse_expr
import tempfile
import docx
import easyocr
from pdf2image import convert_from_bytes
from latex2mathml.converter import convert


class MathTextProcessor:
    def __init__(self):
        self.reader = easyocr.Reader(['vi', 'en'])
        self.custom_config = r'--oem 3 --psm 6 -l vie+eng'
        
    def preprocess_image(self, image):
        img_array = np.array(image)
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        blurred = cv2.GaussianBlur(gray, (3, 3), 0)
        thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(thresh)
        denoised = cv2.fastNlMeansDenoising(enhanced)
        return denoised

    def extract_text_and_math(self, image):
        processed_img = self.preprocess_image(image)
        easy_results = self.reader.readtext(processed_img)
        tess_text = pytesseract.image_to_string(processed_img, config=self.custom_config)
        return self.combine_results(easy_results, tess_text)

    def combine_results(self, easy_results, tess_text):
        combined_text = ""
        math_expressions = []
        for detection in easy_results:
            text = detection[1]
            if self.is_math_expression(text):
                math_expressions.append(text)
            else:
                combined_text += text + "\n"
        math_parts = self.extract_math_expressions(tess_text)
        math_expressions.extend(math_parts)
        return combined_text, math_expressions

    def is_math_expression(self, text):
        math_patterns = [r'[\d+\-*/=(){}[\]]+', r'\bfrac\b|\bsqrt\b', r'∫|∑|√']
        return any(re.search(pattern, text) for pattern in math_patterns)

    def extract_math_expressions(self, text):
        expressions = re.findall(r'[$]{1,2}[^$]+[$]{1,2}', text)
        return list(set(expressions))

    def convert_to_latex(self, expression):
        try:
            expression = self.preprocess_math_expression(expression)
            latex_expr = latex(parse_expr(expression, evaluate=False))
            return f"${latex_expr}$"
        except:
            return expression

    def preprocess_math_expression(self, expr):
        expr = expr.replace('×', '*').replace('÷', '/').replace('^', '**').replace('√', 'sqrt').replace('π', 'pi')
        return expr


def create_results_document(text, math_expressions, latex_expressions):
    doc = docx.Document()
    doc.add_heading('Kết quả chuyển đổi văn bản và công thức toán học', 0)
    doc.add_heading('Văn bản được nhận dạng:', level=1)
    doc.add_paragraph(text)
    doc.add_heading('Các công thức toán học được phát hiện:', level=1)
    for expr in math_expressions:
        doc.add_paragraph(expr)
    doc.add_heading('Mã LaTeX tương ứng:', level=1)
    for latex in latex_expressions:
        doc.add_paragraph(latex)
    return doc


def main():
    st.set_page_config(page_title="Chuyển đổi Văn bản và Công thức Toán học", layout="wide")
    st.title("🔄 Chuyển đổi Văn bản và Công thức Toán học")
    processor = MathTextProcessor()
    uploaded_file = st.file_uploader("Tải lên ảnh hoặc PDF", type=['png', 'jpg', 'jpeg', 'pdf'])

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            images = convert_from_bytes(uploaded_file.read())
            image = images[0]
        else:
            image = Image.open(uploaded_file)

        st.image(image, caption="Ảnh đã tải lên", use_column_width=True)
        if st.button("Bắt đầu chuyển đổi"):
            with st.spinner("Đang xử lý..."):
                text, math_expressions = processor.extract_text_and_math(image)
                latex_expressions = [processor.convert_to_latex(expr) for expr in math_expressions]
                doc = create_results_document(text, math_expressions, latex_expressions)
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        st.download_button("📥 Tải kết quả (DOCX)", f.read(), "ket_qua_chuyen_doi.docx")


if __name__ == "__main__":
    main()

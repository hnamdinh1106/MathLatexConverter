import streamlit as st
import numpy as np
import cv2
from PIL import Image
import pytesseract
import re
from sympy import latex, sympify, parse_expr
import io
import tempfile
import docx
import easyocr
import math
from transformers import pipeline
import pandas as pd
from latex2mathml.converter import convert
from pdf2image import convert_from_bytes
import matplotlib.pyplot as plt

class MathTextProcessor:
    def __init__(self):
        # Khởi tạo EasyOCR cho tiếng Việt và tiếng Anh
        self.reader = easyocr.Reader(['vi', 'en'])
        # Khởi tạo Tesseract với config đặc biệt cho công thức toán
        self.custom_config = r'--oem 3 --psm 6 -l vie+eng'
        
    def preprocess_image(self, image):
        # Nâng cao xử lý ảnh
        img_array = np.array(image)
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        
        # Áp dụng nhiều kỹ thuật tiền xử lý
        blurred = cv2.GaussianBlur(gray, (3, 3), 0)
        thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        
        # Tăng cường độ tương phản
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(thresh)
        
        # Khử nhiễu
        denoised = cv2.fastNlMeansDenoising(enhanced)
        
        return denoised

    def detect_math_regions(self, image):
        # Phát hiện vùng chứa công thức toán học
        gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
        edges = cv2.Canny(gray, 50, 150)
        contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        math_regions = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            if w > 30 and h > 20:  # Lọc các vùng quá nhỏ
                math_regions.append((x, y, w, h))
        
        return math_regions

    def extract_text_and_math(self, image):
        # Xử lý song song văn bản và công thức
        processed_img = self.preprocess_image(image)
        
        # Sử dụng EasyOCR cho văn bản tiếng Việt
        easy_results = self.reader.readtext(processed_img)
        
        # Sử dụng Tesseract cho công thức toán
        tess_text = pytesseract.image_to_string(processed_img, config=self.custom_config)
        
        return self.combine_results(easy_results, tess_text)

    def combine_results(self, easy_results, tess_text):
        combined_text = ""
        math_expressions = []
        
        # Xử lý kết quả EasyOCR
        for detection in easy_results:
            text = detection[1]
            if self.is_math_expression(text):
                math_expressions.append(text)
            else:
                combined_text += text + "\n"
        
        # Xử lý kết quả Tesseract
        math_parts = self.extract_math_expressions(tess_text)
        math_expressions.extend(math_parts)
        
        return combined_text, math_expressions

    def is_math_expression(self, text):
        # Mở rộng nhận dạng biểu thức toán học
        math_patterns = [
            r'[\d+\-*/=(){}[\]]+',
            r'[a-zA-Z]+\s*=\s*.*',
            r'\b[xyz]\b',
            r'\b\d+\s*[+\-*/]\s*\d+\b',
            r'\bsin\b|\bcos\b|\btan\b|\blog\b|\bln\b',
            r'\bπ\b|\balpha\b|\bbeta\b|\bgamma\b',
            r'\bsum\b|\bint\b|\blim\b',
            r'\bfrac\b|\bsqrt\b',
            r'\^',
            r'∫|∑|∏|√|∞|≠|≤|≥|±|∈|∀|∃'
        ]
        
        return any(re.search(pattern, text) for pattern in math_patterns)

    def extract_math_expressions(self, text):
        # Cải thiện trích xuất công thức toán học
        expressions = []
        lines = text.split('\n')
        
        for line in lines:
            # Tìm các công thức toán học trong dòng
            math_parts = re.findall(r'[$]{1,2}[^$]+[$]{1,2}|\\\([^\\]+\\\)|\b\d+[\d\s+\-*/()=><≤≥]+\d+\b', line)
            expressions.extend(math_parts)
            
            # Tìm các biểu thức đại số
            algebra = re.findall(r'[a-zA-Z]+\s*[=]\s*[^.;]*', line)
            expressions.extend(algebra)
            
            # Tìm các phương trình
            equations = re.findall(r'[a-zA-Z\d]+\s*[+\-*/=]\s*[^.;]*', line)
            expressions.extend(equations)
        
        return list(set(expressions))  # Loại bỏ trùng lặp

    def convert_to_latex(self, expression):
        try:
            # Xử lý các trường hợp đặc biệt
            expression = self.preprocess_math_expression(expression)
            
            # Chuyển đổi sang LaTeX
            latex_expr = latex(parse_expr(expression, evaluate=False))
            return f"${latex_expr}$"
        except:
            # Nếu không thể chuyển tự động, áp dụng các quy tắc chuyển đổi cơ bản
            return self.basic_math_to_latex(expression)

    def preprocess_math_expression(self, expr):
        # Chuẩn hóa biểu thức toán học
        expr = expr.replace('×', '*')
        expr = expr.replace('÷', '/')
        expr = expr.replace('^', '**')
        expr = expr.replace('√', 'sqrt')
        expr = expr.replace('π', 'pi')
        return expr

    def basic_math_to_latex(self, expr):
        # Chuyển đổi cơ bản sang LaTeX
        replacements = {
            'sqrt': '\\sqrt',
            'pi': '\\pi',
            '>=': '\\geq',
            '<=': '\\leq',
            '!=': '\\neq',
            'inf': '\\infty',
            'sum': '\\sum',
            'int': '\\int',
            'alpha': '\\alpha',
            'beta': '\\beta',
            'gamma': '\\gamma'
        }
        
        for old, new in replacements.items():
            expr = expr.replace(old, new)
            
        # Xử lý phân số
        fraction_pattern = r'(\d+)/(\d+)'
        expr = re.sub(fraction_pattern, r'\\frac{\1}{\2}', expr)
        
        return f"${expr}$"

def create_results_document(text, math_expressions, latex_expressions):
    doc = docx.Document()
    
    # Thêm tiêu đề
    doc.add_heading('Kết quả chuyển đổi văn bản và công thức toán học', 0)
    
    # Thêm văn bản gốc
    doc.add_heading('Văn bản được nhận dạng:', level=1)
    doc.add_paragraph(text)
    
    # Thêm công thức toán học
    doc.add_heading('Các công thức toán học được phát hiện:', level=1)
    for expr in math_expressions:
        doc.add_paragraph(expr)
    
    # Thêm mã LaTeX
    doc.add_heading('Mã LaTeX tương ứng:', level=1)
    for latex in latex_expressions:
        doc.add_paragraph(latex)
    
    return doc

def main():
    st.set_page_config(page_title="Chuyển đổi Văn bản và Công thức Toán học", layout="wide")
    
    st.title("🔄 Chuyển đổi Văn bản và Công thức Toán học")
    
    # Khởi tạo processor
    processor = MathTextProcessor()
    
    # Upload ảnh
    uploaded_file = st.file_uploader("Tải lên ảnh chứa văn bản và công thức", type=['png', 'jpg', 'jpeg', 'pdf'])
    
    if uploaded_file is not None:
        # Xử lý PDF nếu cần
        if uploaded_file.type == "application/pdf":
            images = convert_from_bytes(uploaded_file.read())
            image = images[0]
        else:
            image = Image.open(uploaded_file)
        
        # Hiển thị ảnh gốc
        st.image(image, caption="Ảnh đã tải lên", use_column_width=True)
        
        if st.button("Bắt đầu chuyển đổi"):
            with st.spinner("Đang xử lý..."):
                # Trích xuất văn bản và công thức
                text, math_expressions = processor.extract_text_and_math(image)
                
                # Chuyển đổi sang LaTeX
                latex_expressions = [processor.convert_to_latex(expr) for expr in math_expressions]
                
                # Hiển thị kết quả
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("### Văn bản")
                    st.text_area("", text, height=300)
                
                with col2:
                    st.markdown("### Công thức toán học")
                    for expr in math_expressions:
                        st.code(expr)
                
                with col3:
                    st.markdown("### Mã LaTeX")
                    for latex in latex_expressions:
                        st.code(latex)
                
                # Tạo file kết quả
                doc = create_results_document(text, math_expressions, latex_expressions)
                
                # Lưu và tạo nút tải xuống
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        st.download_button(
                            label="📥 Tải kết quả (DOCX)",
                            data=f.read(),
                            file_name="ket_qua_chuyen_doi.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

if __name__ == "__main__":
    main()
